import streamlit as st
import pdf2docx
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import io
import os
import tempfile
import tabula
import pandas as pd

# Streamlit app configuration
st.title("PDF to Word Converter with Table Extraction (Supports Nepali Fonts)")
st.write("Upload a PDF to convert to a Word document with tables preserved and Nepali font support. Image support is disabled due to server limitations.")

# File uploader
uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])

def set_nepali_font(run, text):
    """Apply Nepali font (Mangal) to a run."""
    run.text = text
    run.font.name = 'Mangal'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Mangal')

def add_table_to_doc(doc, table_data):
    """Add a table to the Word document with Nepali font."""
    if not table_data or not isinstance(table_data, list) or not table_data[0]:
        return
    rows = len(table_data)
    cols = len(table_data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            run = cell.paragraphs[0].add_run()
            set_nepali_font(run, str(cell_text) if cell_text else "")

def convert_pdf_to_docx(pdf_file):
    """Convert PDF to Word document, including tables."""
    doc = Document()
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
        temp_pdf.write(pdf_file.read())
        temp_pdf_path = temp_pdf.name
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
        temp_docx_path = temp_docx.name
    
    try:
        # Convert PDF text
        cv = Converter(temp_pdf_path)
        cv.convert(temp_docx_path, start=0, end=None)
        cv.close()
        
        # Load converted DOCX
        temp_doc = Document(temp_docx_path)
        for para in temp_doc.paragraphs:
            new_para = doc.add_paragraph()
            run = new_para.add_run()
            set_nepali_font(run, para.text)
        
        # Extract tables using tabula
        tables = tabula.read_pdf(temp_pdf_path, pages='all', multiple_tables=True)
        for table in tables:
            doc.add_paragraph()  # Space before table
            table_data = table.fillna("").values.tolist()
            add_table_to_doc(doc, table_data)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    finally:
        os.unlink(temp_pdf_path)
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)

if uploaded_file is not None:
    file_type = uploaded_file.type
    st.write(f"Processing {uploaded_file.name}...")
    
    try:
        if file_type == "application/pdf":
            # Process PDF
            docx_file = convert_pdf_to_docx(uploaded_file)
        else:
            st.error("Only PDF files are supported in this version!")
            st.stop()
        
        # Provide download link
        st.success("Conversion successful!")
        st.download_button(
            label="Download Word Document",
            data=docx_file,
            file_name="converted_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")