import streamlit as st
import pdf2docx
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import pytesseract
from PIL import Image
import io
import os
import tempfile
import tabula
import pandas as pd

# Streamlit app configuration
st.title("PDF/Image to Word Converter with Table Extraction (Supports Nepali Fonts)")
st.write("Upload a PDF or scanned image to convert to a Word document with tables preserved and Nepali font support.")

# File uploader
uploaded_file = st.file_uploader("Choose a PDF or image file", type=["pdf", "png", "jpg", "jpeg"])

def set_nepali_font(run, text):
    """Apply Nepali font (Mangal) to a run."""
    run.text = text
    run.font.name = 'Mangal'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Mangal')

def add_table_to_doc(doc, table_data):
    """Add a table to the Word document with Nepali font."""
    if not table_data or not isinstance(table_data, list) or not table_data[0]:
        return
    rows = len(table_data)
    cols = len(table_data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            run = cell.paragraphs[0].add_run()
            set_nepali_font(run, str(cell_text) if cell_text else "")

def convert_pdf_to_docx(pdf_file):
    """Convert PDF to Word document, including tables."""
    doc = Document()
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
        temp_pdf.write(pdf_file.read())
        temp_pdf_path = temp_pdf.name
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
        temp_docx_path = temp_docx.name
    
    try:
        # Convert PDF text
        cv = Converter(temp_pdf_path)
        cv.convert(temp_docx_path, start=0, end=None)
        cv.close()
        
        # Load converted DOCX
        temp_doc = Document(temp_docx_path)
        for para in temp_doc.paragraphs:
            new_para = doc.add_paragraph()
            run = new_para.add_run()
            set_nepali_font(run, para.text)
        
        # Extract tables using tabula
        tables = tabula.read_pdf(temp_pdf_path, pages='all', multiple_tables=True)
        for table in tables:
            doc.add_paragraph()  # Space before table
            table_data = table.fillna("").values.tolist()
            add_table_to_doc(doc, table_data)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    finally:
        os.unlink(temp_pdf_path)
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)

def convert_image_to_docx(image_file):
    """Convert scanned image to Word document with OCR, including tables."""
    doc = Document()
    image = Image.open(image_file)
    
    # Perform OCR with Nepali language support
    text = pytesseract.image_to_string(image, lang='nep+eng')
    
    # Add extracted text
    para = doc.add_paragraph()
    set_nepali_font(para.add_run(), text)
    
    # Attempt table detection using pytesseract's data output
    data = pytesseract.image_to_data(image, lang='nep+eng', output_type=pytesseract.Output.DICT)
    table_data = []
    current_row = []
    last_top = -1
    for i, text in enumerate(data['text']):
        if text.strip():
            top = data['top'][i]
            if last_top != -1 and abs(top - last_top) > 10:  # New row
                if current_row:
                    table_data.append(current_row)
                current_row = []
            current_row.append(text)
            last_top = top
    if current_row:
        table_data.append(current_row)
    
    if table_data:
        doc.add_paragraph()  # Space before table
        add_table_to_doc(doc, table_data)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

if uploaded_file is not None:
    file_type = uploaded_file.type
    st.write(f"Processing {uploaded_file.name}...")
    
    try:
        if file_type == "application/pdf":
            # Process PDF
            docx_file = convert_pdf_to_docx(uploaded_file)
        elif file_type in ["image/png", "image/jpeg", "image/jpg"]:
            # Process image
            docx_file = convert_image_to_docx(uploaded_file)
        else:
            st.error("Unsupported file type!")
            st.stop()
        
        # Provide download link
        st.success("Conversion successful!")
        st.download_button(
            label="Download Word Document",
            data=docx_file,
            file_name="converted_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")